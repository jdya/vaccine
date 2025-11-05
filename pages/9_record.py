#pip install streamlit openai python-docx

# ========================================
# 교사용: 생기부 작성 도우미 (고도화)
# ========================================

import io
import streamlit as st
from docx import Document

from utils.session_manager import require_teacher_or_admin, get_current_user, add_message, get_messages, clear_messages, set_current_page
from ai.deepseek_handler import generate_chat_response, stream_chat_response
from database.supabase_manager import get_category_by_name, save_conversation
from utils.helpers import render_auth_modals, render_sidebar_auth_controls, render_sidebar_navigation

st.set_page_config(page_title="📝 생기부", page_icon="📝", layout="wide")
with st.sidebar:
    render_sidebar_navigation()
    render_sidebar_auth_controls()
render_auth_modals()
require_teacher_or_admin()
user = get_current_user()
set_current_page('saeungbu')

st.title("📝 생기부 작성 도우미 (교사용)")

style = st.selectbox("스타일 선택", ["간결형","구체형","발전 중심형"], index=0)
placeholder = {
    "간결형": "관찰 요약을 입력하세요 (예: 수학 흥미 높고 문제 집중도 좋음, 동료 도움)",
    "구체형": "구체 사례를 포함해 입력하세요 (예: 대회 입상, 멘토 역할 등)",
    "발전 중심형": "성장 과정을 중심으로 입력하세요 (예: 최근 변화, 노력 사항)"
}[style]

with st.container(border=True):
    st.subheader("대화형 작성")
    last_ai = None
    for m in get_messages():
        st.chat_message(m['role']).write(m['content'])
        if m['role'] == 'assistant':
            last_ai = m['content']

    text = st.chat_input(placeholder)
    if text:
        # 프롬프트에 스타일 힌트 추가
        combined = f"[스타일:{style}]\n{text}"
        add_message('user', combined)
        assistant_box = st.chat_message("assistant")
        placeholder_box = assistant_box.empty()
        full_text = ""
        for chunk in stream_chat_response(
            category='saeungbu',
            grade=user.get('grade'),
            is_teacher=True,
            conversation_messages=[{"role": m['role'], "content": m['content']} for m in get_messages()],
            temperature=st.session_state.get('ai_temperature', 0.7),
            max_tokens=st.session_state.get('ai_max_tokens', 800)
        ):
            full_text += chunk or ""
            placeholder_box.markdown(full_text)

        if full_text.strip():
            add_message('assistant', full_text)
            cat_row = get_category_by_name('saeungbu')
            if cat_row:
                save_conversation(user['id'], cat_row['id'], combined, full_text, session_id='teacher_saeungbu', is_private=False)
            st.rerun()

col1, col2, col3 = st.columns(3)
with col1:
    if st.button("대화 지우기"):
        clear_messages()
        st.rerun()
with col2:
    if last_ai:
        st.download_button("🗂️ 문장 다운로드(.txt)", data=last_ai, file_name="saeungbu.txt")
with col3:
    if last_ai:
        # Word 문서 생성 후 메모리 버퍼로 제공
        buf = io.BytesIO()
        doc = Document()
        doc.add_heading('생기부 문장', level=1)
        for line in last_ai.splitlines():
            doc.add_paragraph(line)
        doc.save(buf)
        buf.seek(0)
        st.download_button("📝 Word로 내보내기(.docx)", data=buf, file_name="saeungbu.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
